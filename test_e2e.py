"""End-to-end test: exercises the real Flask routes against the real MongoDB
(no mocks) to catch regressions across upload, parsing, row editing,
undo/redo, verification, comments (including anchored), and delete.

Skips the AI review / TTS pipeline stages (no API key / voice pack assumed
in CI) by building the episode directly in the database, the same way the
parse stage would leave it, then testing every route that operates on
already-parsed data. Run with: python test_e2e.py
"""
import io
import os
import unittest

os.environ.setdefault("MONGO_DB_NAME", "audiobook_review_test")

from dotenv import load_dotenv
load_dotenv()
os.environ["MONGO_DB_NAME"] = "audiobook_review_test"

from bson import ObjectId
import time

from docx import Document

import app as app_module
from core import db
from core.jobs import is_running
from config.languages import has_tts_backend


def _build_silent_wav(duration_seconds: float = 1.0, sample_rate: int = 8000) -> bytes:
    """Tiny valid mono 16-bit silent wav, for reviewer-audio-upload tests without a fixture file."""
    import wave
    buffer = io.BytesIO()
    with wave.open(buffer, "wb") as wav:
        wav.setnchannels(1)
        wav.setsampwidth(2)
        wav.setframerate(sample_rate)
        wav.writeframes(b"\x00\x00" * int(sample_rate * duration_seconds))
    return buffer.getvalue()


def build_docx(title: str, chapter_title: str, lines: list[tuple[str, str]]) -> bytes:
    doc = Document()
    doc.add_paragraph(f"Title: {title}")
    doc.add_paragraph("Chapter 1 - " + chapter_title)
    doc.add_paragraph("[")
    for speaker, text in lines:
        doc.add_paragraph(f'"{speaker}: {text}",')
    doc.add_paragraph("]")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class EndToEndTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        app_module.app.config["TESTING"] = True
        cls.client = app_module.app.test_client()
        cls.episode_ids = []

    @classmethod
    def tearDownClass(cls):
        for episode_id in cls.episode_ids:
            db.episodes_collection().delete_one({"_id": ObjectId(episode_id)})
        db.get_db().client.close()

    def _track(self, episode_id):
        self.episode_ids.append(episode_id)
        return episode_id

    def test_landing_page_loads(self):
        resp = self.client.get("/")
        self.assertEqual(resp.status_code, 200)

    def test_upload_creates_episode_and_returns_immediately(self):
        english = build_docx("Test Episode", "Chapter One", [
            ("Narrator", "The sun rose over the hills."),
            ("Mia", "What a beautiful morning."),
        ])
        italian = build_docx("Episodio di Prova", "Capitolo Uno", [
            ("Narrator", "Il sole sorse sulle colline."),
            ("Mia", "Che bella mattina."),
        ])
        resp = self.client.post("/lang/it/new", data={
            "title": "E2E Test Episode",
            "english_docx": (io.BytesIO(english), "english.docx"),
            "translated_docx": (io.BytesIO(italian), "translated.docx"),
        }, content_type="multipart/form-data", follow_redirects=False)
        self.assertEqual(resp.status_code, 302)
        episode_id = resp.location.rstrip("/").split("/")[-1]
        self._track(episode_id)

        episode = db.get_episode(episode_id)
        self.assertIsNotNone(episode)
        self.assertIn(episode["status"], ("uploaded", "parsing", "translating_titles",
                                           "reviewing", "finding_difficult_words", "tts", "failed", "done"))

        for _ in range(50):
            if not is_running(episode_id):
                break
            time.sleep(0.1)

    def _make_parsed_episode(self, target_lang="it", target_lang_name="Italian"):
        episode_id = db.create_episode(title="Direct Test Episode",
                                        target_lang=target_lang, target_lang_name=target_lang_name)
        self._track(episode_id)
        db.set_episode_chapters(episode_id, [{
            "chapter_number": 1, "title": "Chapter One", "translated_title": "Capitolo Uno",
            "title_audio_path": None,
            "title_reviewer_text": "Capitolo Uno",
            "title_reviewer_history": [], "title_reviewer_history_index": -1,
            "title_comments": {t: [] for t in db.TITLE_COMMENT_TARGETS},
            "title_audio_status": "pending", "title_audio_generated_from_text": None,
            "rows": [
                {
                    "sr_no": 1, "speaker": "Narrator",
                    "english": "The sun rose over the hills.",
                    "translated": "Il sole sorse sulle colline.",
                    "reviewer_text": "Il sole sorse sulle colline.",
                    "reviewer_history": [], "reviewer_history_index": -1,
                    "reviewer_complete": False,
                    "comments": {t: [] for t in db.COMMENT_TARGETS},
                    "difficult_words": None, "review_comment": "Looks accurate.",
                    "review_flag": "ok", "audio_path": None, "audio_status": "done",
                    "human_verified": False,
                },
                {
                    "sr_no": 2, "speaker": "Mia",
                    "english": "What a beautiful morning.",
                    "translated": "Che bella mattina.",
                    "reviewer_text": "Che bella mattina.",
                    "reviewer_history": [], "reviewer_history_index": -1,
                    "reviewer_complete": False,
                    "comments": {t: [] for t in db.COMMENT_TARGETS},
                    "difficult_words": None, "review_comment": None,
                    "review_flag": None, "audio_path": None, "audio_status": "pending",
                    "human_verified": False,
                },
            ],
        }])
        db.update_episode(
            episode_id,
            translated_title="Episodio di Prova",
            title_reviewer_text="Episodio di Prova",
            title_reviewer_history=[], title_reviewer_history_index=-1,
            title_comments={t: [] for t in db.TITLE_COMMENT_TARGETS},
            title_audio_status="pending", title_audio_generated_from_text=None,
        )
        db.set_episode_status(episode_id, "done")
        return episode_id

    def test_episode_view_renders(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.get(f"/episode/{episode_id}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b"Il sole sorse sulle colline", resp.data)

    def test_verify_toggle_round_trip(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1",
                                 data={"human_verified": "true"},
                                 headers={"X-Requested-With": "XMLHttpRequest"})
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["verified_rows"], 1)
        self.assertEqual(data["total_rows"], 2)

        row = db.get_row(episode_id, 1)
        self.assertTrue(row["human_verified"])

    def test_reviewer_text_autosave_and_undo_redo(self):
        episode_id = self._make_parsed_episode()

        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-text",
                                 data={"text": "Il sole e sorto sulle colline."})
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertEqual(data["text"], "Il sole e sorto sulle colline.")
        self.assertFalse(data["can_undo"])
        self.assertFalse(data["can_redo"])

        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-text",
                                 data={"text": "Il sole sorse ancora sulle colline."})
        data = resp.get_json()
        self.assertTrue(data["can_undo"])
        self.assertFalse(data["can_redo"])

        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-text/undo")
        self.assertEqual(resp.get_json()["text"], "Il sole e sorto sulle colline.")

        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-text/redo")
        self.assertEqual(resp.get_json()["text"], "Il sole sorse ancora sulle colline.")

    def test_reviewer_text_missing_row_returns_404(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/999/reviewer-text",
                                 data={"text": "no such row"})
        self.assertEqual(resp.status_code, 404)

    def test_complete_toggle(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1/complete", data={"complete": "true"})
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.get_json()["complete"])
        self.assertTrue(db.get_row(episode_id, 1)["reviewer_complete"])

    def test_general_comment_round_trip(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1/comments/reviewer_edit",
                                 data={"text": "Consider a different phrasing", "author": "Reviewer"})
        self.assertEqual(resp.status_code, 200)
        comment = resp.get_json()["comment"]
        self.assertEqual(comment["text"], "Consider a different phrasing")
        self.assertIsNone(comment["anchor"])

        resp = self.client.post(
            f"/episode/{episode_id}/row/1/comments/reviewer_edit/{comment['id']}/reply",
            data={"text": "Good point", "author": "Second Reviewer"})
        self.assertEqual(resp.status_code, 200)

        resp = self.client.post(
            f"/episode/{episode_id}/row/1/comments/reviewer_edit/{comment['id']}/resolve",
            data={"resolved": "true"})
        self.assertTrue(resp.get_json()["resolved"])

        resp = self.client.post(
            f"/episode/{episode_id}/row/1/comments/reviewer_edit/{comment['id']}/delete")
        self.assertTrue(resp.get_json()["ok"])

        row = db.get_row(episode_id, 1)
        self.assertEqual(row["comments"]["reviewer_edit"], [])

    def test_anchored_comment_on_english_target(self):
        episode_id = self._make_parsed_episode()
        text = "The sun rose over the hills."
        quote = "sun rose"
        start = text.index(quote)
        end = start + len(quote)

        resp = self.client.post(f"/episode/{episode_id}/row/1/comments/english", data={
            "text": "Nice imagery here", "author": "Reviewer",
            "anchor_quote": quote, "anchor_start": str(start), "anchor_end": str(end),
        })
        self.assertEqual(resp.status_code, 200)
        comment = resp.get_json()["comment"]
        self.assertEqual(comment["anchor"]["quote"], quote)
        self.assertEqual(comment["anchor"]["start"], start)
        self.assertEqual(comment["anchor"]["end"], end)

    def test_anchored_comment_rejected_on_non_anchorable_target(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1/comments/reviewer_edit", data={
            "text": "Should not anchor", "author": "Reviewer",
            "anchor_quote": "sun", "anchor_start": "0", "anchor_end": "3",
        })
        self.assertEqual(resp.status_code, 200)
        self.assertIsNone(resp.get_json()["comment"]["anchor"])

    def test_anchor_with_invalid_range_is_rejected(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1/comments/english", data={
            "text": "Bad range", "author": "Reviewer",
            "anchor_quote": "sun", "anchor_start": "10", "anchor_end": "2",
        })
        self.assertEqual(resp.status_code, 200)
        self.assertIsNone(resp.get_json()["comment"]["anchor"])

    def test_invalid_comment_target_is_404(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/row/1/comments/not_a_real_target",
                                 data={"text": "x", "author": "Reviewer"})
        self.assertEqual(resp.status_code, 404)

    def test_status_endpoint_reports_progress(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.get(f"/episode/{episode_id}/status")
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertEqual(data["status"], "done")
        self.assertEqual(data["total_rows"], 2)
        self.assertFalse(data["stalled"])

    def test_audio_status_endpoint(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.get(f"/episode/{episode_id}/audio-status")
        self.assertEqual(resp.status_code, 200)
        rows = resp.get_json()["rows"]
        self.assertEqual(len(rows), 2)
        self.assertEqual({r["sr_no"] for r in rows}, {1, 2})

    def test_delete_episode(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/delete", follow_redirects=False)
        self.assertEqual(resp.status_code, 302)
        self.assertIsNone(db.get_episode(episode_id))
        self.episode_ids.remove(episode_id)

    def test_regenerate_audio_uses_reviewer_text(self):
        episode_id = self._make_parsed_episode()
        if not has_tts_backend("it"):
            resp = self.client.post(f"/episode/{episode_id}/row/1/regenerate-audio")
            self.assertEqual(resp.status_code, 400)
            return
        self.client.post(f"/episode/{episode_id}/row/1/reviewer-text",
                          data={"text": "Il sole e sorto di nuovo sulle colline."})
        resp = self.client.post(f"/episode/{episode_id}/row/1/regenerate-audio")
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertTrue(data["ok"])
        row = db.get_row(episode_id, 1)
        self.assertEqual(row["reviewer_text"], row["audio_generated_from_text"])

    def test_reviewer_audio_upload_and_delete_round_trip(self):
        episode_id = self._make_parsed_episode()
        wav_bytes = _build_silent_wav()
        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-audio",
                                 data={"audio": (io.BytesIO(wav_bytes), "take1.wav")},
                                 content_type="multipart/form-data")
        self.assertEqual(resp.status_code, 200)
        data = resp.get_json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["audio_source"], "reviewer")
        row = db.get_row(episode_id, 1)
        self.assertIsNotNone(row["reviewer_audio_path"])
        self.assertEqual(row["audio_source"], "reviewer")

        resp = self.client.post(f"/episode/{episode_id}/row/1/reviewer-audio/delete")
        self.assertEqual(resp.status_code, 200)
        row = db.get_row(episode_id, 1)
        self.assertIsNone(row["reviewer_audio_path"])
        self.assertEqual(row["audio_source"], "tts")

    def test_audio_source_toggle(self):
        episode_id = self._make_parsed_episode()
        db.update_row(episode_id, 1, audio_path="row_1.mp3", reviewer_audio_path="row_1_reviewer.mp3",
                       audio_source="tts")
        resp = self.client.post(f"/episode/{episode_id}/row/1/audio-source", data={"source": "reviewer"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(db.get_row(episode_id, 1)["audio_source"], "reviewer")

        resp = self.client.post(f"/episode/{episode_id}/row/1/audio-source", data={"source": "tts"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(db.get_row(episode_id, 1)["audio_source"], "tts")

    def test_stale_audio_badge_reflects_text_edit(self):
        episode_id = self._make_parsed_episode()
        db.update_row(episode_id, 1, audio_path="row_1.mp3",
                       audio_generated_from_text="Il sole sorse sulle colline.")
        self.client.post(f"/episode/{episode_id}/row/1/reviewer-text", data={"text": "Testo modificato."})
        row = db.get_row(episode_id, 1)
        self.assertNotEqual(row["reviewer_text"], row["audio_generated_from_text"])

    def test_title_reviewer_text_autosave_and_undo_redo(self):
        episode_id = self._make_parsed_episode()

        resp = self.client.post(f"/episode/{episode_id}/title/reviewer-text",
                                 data={"text": "Episodio Modificato"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.get_json()["text"], "Episodio Modificato")

        resp = self.client.post(f"/episode/{episode_id}/title/reviewer-text",
                                 data={"text": "Episodio Modificato Ancora"})
        self.assertTrue(resp.get_json()["can_undo"])

        resp = self.client.post(f"/episode/{episode_id}/title/reviewer-text/undo")
        self.assertEqual(resp.get_json()["text"], "Episodio Modificato")

        resp = self.client.post(f"/episode/{episode_id}/title/reviewer-text/redo")
        self.assertEqual(resp.get_json()["text"], "Episodio Modificato Ancora")

        resp = self.client.post(f"/episode/{episode_id}/chapter/1/title/reviewer-text",
                                 data={"text": "Capitolo Modificato"})
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.get_json()["text"], "Capitolo Modificato")

    def test_chapter_title_routes_missing_chapter_returns_404(self):
        episode_id = self._make_parsed_episode()
        for resp in (
            self.client.post(f"/episode/{episode_id}/chapter/999/title/reviewer-text", data={"text": "x"}),
            self.client.post(f"/episode/{episode_id}/chapter/999/title/reviewer-text/undo"),
            self.client.post(f"/episode/{episode_id}/chapter/999/title/reviewer-text/redo"),
            self.client.post(f"/episode/{episode_id}/chapter/999/title/regenerate-audio"),
        ):
            self.assertEqual(resp.status_code, 404)

    def test_title_comment_round_trip_and_anchored(self):
        episode_id = self._make_parsed_episode()
        text = "Direct Test Episode"
        quote = "Test Episode"
        start = text.index(quote)
        end = start + len(quote)

        resp = self.client.post(f"/episode/{episode_id}/title/comments/english", data={
            "text": "Nice title", "author": "Reviewer",
            "anchor_quote": quote, "anchor_start": str(start), "anchor_end": str(end),
        })
        self.assertEqual(resp.status_code, 200)
        comment = resp.get_json()["comment"]
        self.assertEqual(comment["anchor"]["quote"], quote)
        episode = db.get_episode(episode_id)
        self.assertEqual(len(episode["title_comments"]["english"]), 1)

        resp = self.client.post(f"/episode/{episode_id}/chapter/1/title/comments/reviewer_edit",
                                 data={"text": "Chapter title note", "author": "Reviewer"})
        self.assertEqual(resp.status_code, 200)
        chapter_comment = resp.get_json()["comment"]
        self.assertIsNone(chapter_comment["anchor"])
        episode = db.get_episode(episode_id)
        self.assertEqual(len(episode["chapters"][0]["title_comments"]["reviewer_edit"]), 1)

    def test_title_comment_invalid_target_is_404(self):
        episode_id = self._make_parsed_episode()
        resp = self.client.post(f"/episode/{episode_id}/title/comments/not_a_real_target",
                                 data={"text": "x", "author": "Reviewer"})
        self.assertEqual(resp.status_code, 404)

        resp = self.client.post(f"/episode/{episode_id}/title/comments/audio",
                                 data={"text": "x", "author": "Reviewer"})
        self.assertEqual(resp.status_code, 404)


if __name__ == "__main__":
    unittest.main(verbosity=2)
