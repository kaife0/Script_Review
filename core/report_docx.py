"""Reviewed-script docx builder — reproduces the input script's format (Title/Theme/Chapter
blocks with quoted "Speaker: text" dialogue lines) but with the reviewer's final text."""
from docx import Document


def build_docx_report(episode: dict) -> Document:
    doc = Document()

    title = episode.get("title_reviewer_text") or episode["title"]
    doc.add_paragraph(f"Title: {title}")
    doc.add_paragraph(f"Theme: {episode.get('theme', '')}")
    doc.add_paragraph("")

    for chapter in episode["chapters"]:
        chapter_title = chapter.get("title_reviewer_text") or chapter["title"]
        doc.add_paragraph(f"Chapter {chapter['chapter_number']}: {chapter_title} [")
        for row in chapter["rows"]:
            text = row.get("reviewer_text") or row.get("translated") or ""
            doc.add_paragraph(f'"{row["speaker"]}: {text}",')
        doc.add_paragraph("]")
        doc.add_paragraph("")

    return doc
